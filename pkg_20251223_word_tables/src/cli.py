from docx import Document
import pandas as pd
import re
from pathlib import Path
from datetime import datetime
from zoneinfo import ZoneInfo
from zipfile import ZipFile
import tempfile
import subprocess
import click
from rich.console import Console
from rich.prompt import Prompt, Confirm
from rich.progress import track
import shutil

from .name_utils import (
    unify_first_last,
    match_csv_docx_names,
)
from ._vars import (
    KTP_FIRST_NAME_COL,
    KTP_LAST_NAME_COL,
    KTP_FIRST_NAME_ORIG_COLNAME_COL,
    KTP_LAST_NAME_ORIG_COLNAME_COL,
    DRAW_LABEL,
    RIGHT_NAME_COL,
)

console = Console()

PACKAGE_ROOT = Path(__file__).parent
PANDOC_REFERENCE_DOCX_PATH = PACKAGE_ROOT / "resources" / "pandoc-custom-reference.docx"

TOTAL_DRAWS = 310  # e.g., as of 2025-12-23 (including pilot)

INTRODUCTION = """## Introduction
**Draw number** is the sequential order in which rows were sampled from HCR tables.

Name is displayed as **Last Name, First Name**.

Last modified (introduction): December 23, 2025

Date of report: {}
"""

def find_files_by_extension(directory: Path, extension: str, recursive: bool = False) -> list[Path]:
    """Find all files with given extension in directory."""
    pattern = f"*.{extension}"
    if recursive:
        return list(directory.rglob(pattern))
    else:
        return list(directory.glob(pattern))
    
def validate_csv_headers(csv_files: list[Path]) -> bool:
    """Validate that all CSV files have the same column names."""
    if not csv_files:
        return False
    
    first_df = pd.read_csv(csv_files[0], nrows=0)
    expected_cols = set(first_df.columns)
    
    for csv_path in csv_files[1:]:
        df = pd.read_csv(csv_path, nrows=0)
        if set(df.columns) != expected_cols:
            console.print(f"[red]Column mismatch in {csv_path.name}[/red]")
            console.print(f"Expected: {sorted(expected_cols)}")
            console.print(f"Got: {sorted(df.columns)}")
            return False
    
    return True

def get_cell_text_with_format(cell):
    texts = []
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            # mark formatting in some way, e.g., HTML-like
            txt = run.text
            if run.bold:
                txt = f"**{txt}**"
            if run.italic:
                txt = f"_{txt}_"
            if run.font.subscript:
                txt = f"~{txt}~"
            if run.font.superscript:
                txt = f"^{txt}^"
            texts.append(txt)
    return ''.join(texts)

def parse_docx_standard(docx_path: Path) -> list[pd.DataFrame]:
    """Parse DOCX using standard python-docx."""
    doc = Document(docx_path)

    dfs = []
    for t in doc.tables:
        rows = []
        for i, row in enumerate(t.rows):
            # Use plain text for the first row, formatted text for the rest
            if i == 0:
                rows.append([cell.text.strip() for cell in row.cells])
            else:
                rows.append([get_cell_text_with_format(cell) for cell in row.cells])

        df = pd.DataFrame(rows)

        df.columns = df.iloc[0]  # first row as header (plain text)
        df = df.iloc[1:].reset_index(drop=True)

        dfs.append(df)
    
    return dfs

def process_documents(docx_dir: Path, csv_dir: Path, recursive: bool, 
                     output_dir: Path, output_format: str):
    """Main processing logic."""
    # Find all DOCX files
    docx_files = find_files_by_extension(docx_dir, "docx", recursive)
    if not docx_files:
        console.print("[red]No DOCX files found in specified directory.[/red]")
        return
    
    console.print(f"[green]Found {len(docx_files)} DOCX file(s)[/green]")
    
    # Find and validate CSV files
    csv_files = find_files_by_extension(csv_dir, "csv", recursive)
    if not csv_files:
        console.print("[red]No CSV files found in specified directory.[/red]")
        return
    
    console.print(f"[green]Found {len(csv_files)} CSV file(s)[/green]")
    
    if not validate_csv_headers(csv_files):
        console.print("[red]CSV files have different headers. Aborting.[/red]")
        return
    
    # Parse DOCX files
    all_dfs = []
    parse_func = parse_docx_standard
    
    for docx_path in track(docx_files, description="Parsing DOCX files..."):
        dfs = parse_func(docx_path)
        all_dfs.extend(dfs)
    
    # load and combine all CSV files
    csv_df = pd.concat([pd.read_csv(csv_path) for csv_path in csv_files], ignore_index=True)

    # docx tables
    docx_df = pd.concat(all_dfs, ignore_index=True)

    # csv join key
    unified_names = csv_df.apply(unify_first_last, axis=1, result_type='expand')
    csv_df[KTP_FIRST_NAME_COL] = unified_names[0].apply(lambda x: x[KTP_FIRST_NAME_COL])
    csv_df[KTP_LAST_NAME_COL] = unified_names[1].apply(lambda x: x[KTP_LAST_NAME_COL])

    docx_df.to_csv('tmp/docx_df.csv') # debug
    
    # Match names and get DOCX indices
    docx_indices = match_csv_docx_names(
        csv_df[KTP_FIRST_NAME_COL],
        csv_df[KTP_LAST_NAME_COL],
        docx_df[RIGHT_NAME_COL],
    )

    # Join using the matched indices
    csv_df['_docx_idx'] = docx_indices
    docx_df['_docx_idx'] = docx_df.index
    joined_df = csv_df.merge(docx_df, on='_docx_idx', how='left')

    # optional cleanup
    joined_df = joined_df.drop(columns=['_docx_idx'])
    joined_df.rename(  # normalize header row for docx table
        columns=lambda col: (
            re.sub(
                pattern=r'_+',  # otherwise consecutive underscores break markdown
                repl='_',
                string=(
                    col if re.match(
                        pattern=r'^[\w_]+\.',  # e.g., starts with `ktp.`, `hcr.` etc.
                        string=str(col),
                    )
                    else 'ktp.table_1_' + re.sub(
                        pattern=r'\s',
                        repl='_',
                        string=re.sub(r'[^\w\s]', '_', str(col).lower()),
                    )
                )
            )
        ),
        inplace=True,
    )

    # Create markdown cards for each row
    cards = {}
    intro = INTRODUCTION.format(
        today := (
            datetime.now(ZoneInfo('America/Toronto'))
            .strftime('%B %d, %Y')
        )
    ) + "\n\n"  # will merge later so as not to spoil cards
    for _, row in joined_df.iterrows():
        draw_number = row.pop(DRAW_LABEL)
        card = (
            f"### Draw #{draw_number} of {TOTAL_DRAWS}: {row.get(KTP_LAST_NAME_COL)}, {row.get(KTP_FIRST_NAME_COL)}\n"
            f"Fun fact: the last name came from `{row.get(KTP_LAST_NAME_ORIG_COLNAME_COL)}` and the first name – from `{row.get(KTP_FIRST_NAME_ORIG_COLNAME_COL)}` in the originating HCR list."
        )
        # docx_filename = re.sub(r'\s+', '_', re.sub(r'[^A-Za-z0-9\s]+', '', card)).strip('_')
        # Less verbose version
        minified_card = f"{draw_number}: {row.get(KTP_FIRST_NAME_COL)} {row.get(KTP_LAST_NAME_COL)}\n"
        docx_filename = re.sub(r'\s+', '_', re.sub(r'[^A-Za-z0-9\s]+', '', minified_card)).strip('_')
        for col, val in row.items():
            if '\n' in str(val):
                # treat as code block
                card += f"**{col}**:\n\n{str(val).replace('\n','\n\n')}\n\n"
            else:
                card += f"**{col}**: {str(val)}\n\n"
        cards[docx_filename] = card

    # Create output directory if it doesn't exist
    output_dir.mkdir(parents=True, exist_ok=True)
    
    # Save as standalone files and zip them
    zip_path = output_dir / f"{csv_dir.name}_combined_cards.zip"
    
    # Save based on output format
    if output_format == "txt":
        with tempfile.TemporaryDirectory() as tmpdir:
            txt_paths = []
            for filename, card in track(list(cards.items()), description="Creating Markdown (*.txt) files..."):
                txt_path = Path(tmpdir) / f"{filename}.txt"
                txt_path.write_text(intro + card, encoding="utf-8")
                txt_paths.append(txt_path)
            # Zip them all
            with ZipFile(zip_path, "w") as zipf:
                for path in txt_paths:
                    zipf.write(path, arcname=path.name)
        console.print(f"[green]Saved Markdown (*.txt) files to: {zip_path}[/green]")
        
    elif output_format == "docx":
        with tempfile.TemporaryDirectory() as tmpdir:
            tmp_ref_path = Path(tmpdir) / Path(PANDOC_REFERENCE_DOCX_PATH).name
            shutil.copy(PANDOC_REFERENCE_DOCX_PATH, tmp_ref_path)
            docx_paths: list[Path] = []
            for filename, card in track(list(cards.items()), description="Creating DOCX files..."):
                md_path = Path(tmpdir) / f"{filename}.md"
                docx_path = Path(tmpdir) / f"{filename}.docx"
                md_path.write_text(intro + card, encoding="utf-8")
                subprocess.run([
                    "pandoc",
                    str(md_path),
                    "-o",
                    str(docx_path),
                    "--reference-doc",
                    str(tmp_ref_path),
                ], check=True)
                docx_paths.append(docx_path)
            with ZipFile(zip_path, "w") as zipf:
                for path in docx_paths:
                    zipf.write(path, arcname=path.name)
        console.print(f"[green]Saved DOCX files to: {zip_path}[/green]")

@click.group(invoke_without_command=True)
@click.pass_context
def cli(ctx):
    """Document enrichment CLI tool."""
    if ctx.invoked_subcommand is None:
        ctx.invoke(interactive)

@cli.command()
def interactive():
    """Run in interactive mode with rich prompts."""
    console.print("[bold blue]Document Enrichment Tool - Interactive Mode[/bold blue]\n")
    
    # Get inputs
    docx_dir = Path(Prompt.ask("Path to directory containing DOCX files"))
    while not docx_dir.exists() or not docx_dir.is_dir():
        console.print("[red]Invalid directory path.[/red]")
        docx_dir = Path(Prompt.ask("Path to directory containing DOCX files"))
    
    recursive = Confirm.ask("Search recursively for DOCX files?", default=False)
    
    csv_dir = Path(Prompt.ask("Path to directory containing CSV files"))
    while not csv_dir.exists() or not csv_dir.is_dir():
        console.print("[red]Invalid directory path.[/red]")
        csv_dir = Path(Prompt.ask("Path to directory containing CSV files"))
    
    output_dir = Path(Prompt.ask("Output directory", default="./output"))
    
    output_format = Prompt.ask(
        "Output format",
        choices=["txt", "docx"],
        default="txt"
    )
    
    # Process
    process_documents(docx_dir, csv_dir, recursive, output_dir, output_format)

@cli.command()
@click.argument('docx_dir', type=click.Path(exists=True, file_okay=False, path_type=Path))
@click.argument('csv_dir', type=click.Path(exists=True, file_okay=False, path_type=Path))
@click.option('-r', '--recursive', is_flag=True, help='Search recursively for DOCX and CSV files')
@click.option('--output-dir', type=click.Path(path_type=Path), default='./output', 
              help='Output directory (default: ./output)')
@click.option('--output-format', type=click.Choice(['txt', 'docx']), default='txt',
              help='Output format: txt (markdown) or docx (default: txt)')
def process(docx_dir: Path, csv_dir: Path, recursive: bool, output_dir: Path, 
           output_format: str):
    """Process DOCX files and enrich with CSV data.
    
    DOCX_DIR: Directory containing DOCX files
    CSV_DIR: Directory containing CSV files (must have matching headers)
    """
    
    process_documents(docx_dir, csv_dir, recursive, output_dir, output_format)


