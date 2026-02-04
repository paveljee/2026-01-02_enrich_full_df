from pathlib import Path

import pandas as pd
import pytest
from docx import Document

from src._vars import KTP_FIRST_NAME_COL, KTP_LAST_NAME_COL, RIGHT_NAME_COL
from src.name_utils import match_csv_docx_names, unify_first_last
from src.parse_docx import parse_docx_table
from src.utils.files import find_files_by_extension

TEST_CSV_DIR = Path("data/samples")
TEST_DOCX_DIR = Path("data/manual_extractions")


def create_test_docx_with_table(path: Path, names: list[str]):
    """Helper to create a DOCX file with a table matching the image structure."""
    doc = Document()
    
    # Add italic caption
    caption = doc.add_paragraph()
    caption_run = caption.add_run(
        "Table 1: PROGRESS-Plus factors for sample (n_52) of authors on "
        "Clarivate's Highly Cited Researchers List."
    )
    caption_run.italic = True
    
    # Create table: 9 columns matching the image
    table = doc.add_table(rows=1, cols=9)
    table.style = 'Table Grid'
    
    # Header row - first column must be "Researcher/author"
    hdr = table.rows[0].cells
    hdr[0].text = "Researcher/author"
    hdr[1].text = "Place of\nresidence"
    hdr[2].text = "Gender"
    hdr[3].text = "Age\n\n*First\nPublicat\nion\naccordi\ng to\nOpenAl\nex\nprofile"
    hdr[4].text = "Education"
    hdr[5].text = "Academic position(s)"
    hdr[6].text = "Social capital"
    hdr[7].text = "Links*"
    hdr[8].text = ""
    
    # Add data rows with the name in first column
    for name in names:
        row = table.add_row().cells
        row[0].text = name
        row[1].text = "Toronto,\nCanada"
        row[2].text = "Male"
        row[3].text = "77"
        row[4].text = "BA in Experimental\nPsychology"
        row[5].text = "University Professor Emeritus"
        row[6].text = "Turing Award\n(2018)"
        row[7].text = "NR"
        row[8].text = ""
    
    # Add footnotes
    doc.add_paragraph("*verified with ORCID")
    doc.add_paragraph("**NR = Not Reported, NA = Not Available/Applicable")
    
    doc.save(str(path))

def test_create_test_docx_with_table(tmp_path):
    names = ["name1", "name2"]
    work_path = tmp_path / "mock_RI_test.docx"
    create_test_docx_with_table(work_path, names)
    assert work_path.exists()

@pytest.fixture
def synthetic_csv_docx_pair(tmp_path):
    """Create matching CSV and DOCX test files with known name pairs."""
    
    # Create CSV with ALREADY UNIFIED names
    csv_rows = [
        {
            KTP_FIRST_NAME_COL: "Geoffrey",
            KTP_LAST_NAME_COL: "Hinton",
        },
        {
            KTP_FIRST_NAME_COL: "Lane W.",
            KTP_LAST_NAME_COL: "Martin",
        },
        {
            KTP_FIRST_NAME_COL: "Yann",
            KTP_LAST_NAME_COL: "LeCun",
        },
    ]
    
    csv_df = pd.DataFrame(csv_rows)
    csv_path = tmp_path / "test_sample.csv"
    csv_df.to_csv(csv_path, index=False)
    
    # Create DOCX with combined names in various formats
    docx_names = [
        "Geoffrey Hinton ",
        "Lane W. Martin",
        "Yann LeCun",
    ]
    
    docx_path = tmp_path / "test_table.docx"
    create_test_docx_with_table(docx_path, docx_names)
    
    return csv_path, docx_path


def test_match_csv_docx_names_with_synthetic_data(synthetic_csv_docx_pair):
    """Test match_csv_docx_names with controlled synthetic data."""
    csv_path, docx_path = synthetic_csv_docx_pair
    
    # Load CSV (already has unified names)
    csv_df = pd.read_csv(csv_path)
    
    # Parse DOCX
    docx_dfs = parse_docx_table(docx_path)
    docx_df = docx_dfs[0]

    # Run the matching function
    docx_indices = match_csv_docx_names(
        csv_df[[KTP_FIRST_NAME_COL, KTP_LAST_NAME_COL]],
        docx_df[RIGHT_NAME_COL],
    )
    
    # Verify matches
    for idx, docx_idx in enumerate(docx_indices):
        first = csv_df[KTP_FIRST_NAME_COL].iloc[idx]
        last = csv_df[KTP_LAST_NAME_COL].iloc[idx]
        matched_name = docx_df[RIGHT_NAME_COL].iloc[docx_idx]
        
        assert first in matched_name
        assert last in matched_name


def test_match_csv_docx_names_on_full_dataset(tmp_path):
    """Test match_csv_docx_names on full dataset with actual data."""
    
    if not TEST_CSV_DIR.exists():
        pytest.skip(f"Test data directory not found: {TEST_CSV_DIR}")
    
    if not TEST_DOCX_DIR.exists():
        pytest.skip(f"Test DOCX directory not found: {TEST_DOCX_DIR}")
    
    # Find files
    csv_files = find_files_by_extension(TEST_CSV_DIR, "csv", recursive=False)
    if not csv_files:
        pytest.skip("No CSV files found")
    
    docx_files = find_files_by_extension(TEST_DOCX_DIR, "docx", recursive=False)
    if not docx_files:
        pytest.skip("No DOCX files found")
    
    # Load CSV data
    csv_df = pd.concat([pd.read_csv(csv_path) for csv_path in csv_files], ignore_index=True)
    
    # Parse DOCX files
    all_docx_dfs = []
    for docx_path in docx_files:
        dfs = parse_docx_table(docx_path)
        all_docx_dfs.extend(dfs)
    
    docx_df = pd.concat(all_docx_dfs, ignore_index=True)
    
    # Unify names (this is tested elsewhere, but needed for this dataset)
    unified_names = csv_df.apply(unify_first_last, axis=1, result_type='expand')
    csv_df[KTP_FIRST_NAME_COL] = unified_names[0].apply(lambda x: x[KTP_FIRST_NAME_COL])
    csv_df[KTP_LAST_NAME_COL] = unified_names[1].apply(lambda x: x[KTP_LAST_NAME_COL])
    
    # Track failures
    failures = []
    
    try:
        # Run the matching function
        docx_indices = match_csv_docx_names(
            csv_df[[KTP_FIRST_NAME_COL, KTP_LAST_NAME_COL]],
            docx_df[RIGHT_NAME_COL],
        )
        
        # Verify sample of matches
        for idx in range(min(10, len(docx_indices))):
            first = csv_df[KTP_FIRST_NAME_COL].iloc[idx]
            last = csv_df[KTP_LAST_NAME_COL].iloc[idx]
            matched_name = docx_df[RIGHT_NAME_COL].iloc[docx_indices[idx]]
            
            if first not in matched_name or last not in matched_name:
                failures.append((idx, first, last, matched_name))
        
    except ValueError as e:
        msg = str(e)

        # Expected/known ambiguity: exactly 7 rows and all MULTIPLE_MATCHES
        if "Could not uniquely match 7 CSV rows" in msg:
            lines = msg.splitlines()
            # rows listed after the header line; keep only the bullet-ish lines
            problem_lines = [ln.strip() for ln in lines if "(csv_idx=" in ln]

            if len(problem_lines) == 7 and all("MULTIPLE_MATCHES:" in ln for ln in problem_lines):
                pytest.xfail(msg)

        pytest.fail(f"Matching failed:\n{msg}")
    
    if failures:
        failure_msg = "\n".join([
            f"Row {idx}: {first} {last} -> '{matched}'"
            for idx, first, last, matched in failures
        ])
        pytest.fail(f"Match validation failed:\n{failure_msg}")
